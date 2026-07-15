#!/usr/bin/env python3
"""按《优化分析方案.docx》第四部分结果呈现模板，将本版 output 表格整合为 Word。"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = Path("/Users/zengyinghao/Downloads/output")
DEFAULT_DOCX = ROOT.parent / "ALL" / "AICU_优化方案结果汇总.docx"

VAR_CN = {
    "shannon": "Shannon指数",
    "log_crp": "log(CRP)",
    "asa": "ASA分级",
    "anesthesia_duration_min": "麻醉时长(min)",
}


def _font(run, *, east: str = "宋体", size: Pt | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    r = run._element.get_or_add_rPr()
    r.rFonts.set(qn("w:eastAsia"), east)


def _h(doc: Document, text: str, level: int = 1) -> None:
    for run in doc.add_heading(text, level=level).runs:
        _font(run)


def _p(doc: Document, text: str, *, bold: bool = False, size: Pt = Pt(11)) -> None:
    run = doc.add_paragraph().add_run(text)
    _font(run, bold=bold, size=size)


def _note(doc: Document, text: str) -> None:
    run = doc.add_paragraph().add_run(text)
    _font(run, size=Pt(9))
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = "" if val is None or (isinstance(val, float) and pd.isna(val)) else str(val)
    doc.add_paragraph()


def _read(path: Path) -> pd.DataFrame | None:
    if not path.exists():
        return None
    return pd.read_csv(path, encoding="utf-8-sig")


def _first(*dfs: pd.DataFrame | None) -> pd.DataFrame | None:
    for d in dfs:
        if d is not None and not d.empty:
            return d
    return None


def _fmt(x, nd: int = 3) -> str:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return "—"
    try:
        return f"{float(x):.{nd}f}"
    except (TypeError, ValueError):
        return str(x)


def _df_rows(df: pd.DataFrame) -> list[list[str]]:
    return [[("" if pd.isna(v) else str(v)) for v in row] for row in df.astype(object).values.tolist()]


def build(output_dir: Path, docx_path: Path) -> Path:
    output_dir = output_dir.resolve()
    tab = output_dir / "tables"
    fig = output_dir / "figures"
    summary = json.loads((output_dir / "run_summary.json").read_text(encoding="utf-8"))

    t1 = _read(tab / "table1_baseline.csv")
    t1_notes = _read(tab / "table1_footnotes.csv")
    t1_sup = _read(tab / "table1_supplement_diversity.csv")
    perm = _read(fig / "figure3_permanova.csv")
    beta = _read(fig / "figure3_betadisper.csv")
    genus = _read(tab / "table_genus_wilcoxon.csv")
    lefse = _read(fig / "figure4_lefse_results.csv")
    t2 = _read(tab / "table2_model_performance.csv")
    t2_fac = _read(tab / "table2_factorial_delta_auc.csv")
    t2_cont = _read(tab / "table2_continuous_extubation.csv")
    t2_supp = _read(tab / "table2_model_performance_supplementary.csv")
    t3 = _read(tab / "table3_multivariable_or.csv")
    corr = _first(
        _read(fig / "figure5_correlations.csv"),
        _read(tab / "table_shannon_inflammation_corr.csv"),
    )
    med = _read(fig / "figure7_mediation.csv")
    ae = _read(tab / "table_ae_descriptive.csv")
    nlr = _read(tab / "nlr_audit.csv")
    cprime = _read(tab / "model_c_prime_decision.csv")
    sens = _read(output_dir / "sensitivity" / "sensitivity_cohort_summary.csv")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")

    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AICU呼吸道菌群预后研究\n优化方案结果汇总（表格整合）")
    _font(r, east="黑体", size=Pt(16), bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"依据：《优化分析方案.docx》第四部分结果呈现模板\n"
        f"数据：n={summary.get('n_samples')} "
        f"（Early {summary.get('n_early')} / Delayed {summary.get('n_delayed')}），"
        f"{'演示菌群' if summary.get('demo_microbiome') else '真实 DADA2 菌群'}\n"
        f"运行时间：{summary.get('run_time', '—')}\n"
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"来源目录：{output_dir}"
    )
    _font(sr, size=Pt(10))

    # 〇、变量与模型定义（依据《优化分析方案.docx》第3.2(B) 特征精简建议）
    _h(doc, "〇、变量与模型定义说明", 1)
    _p(
        doc,
        "本节按《优化分析方案.docx》第3.2(B)「特征精简建议（避免 n=34 下过拟合）」定义本版所有分析变量与模型，"
        "并说明与旧版《AICU_output表格对照解析报告.docx》的差异，供跨文档对照时避免误读。",
    )

    _h(doc, "0.1 变量定义与口径", 2)
    _table(
        doc,
        ["变量（代码名）", "含义", "口径 / 处理方式", "在方案中的角色"],
        [
            ["ASA分级 (asa)", "美国麻醉医师协会体格分级", "术前评估分级，连续化纳入", "临床基础特征（与 Table 3 对齐）"],
            ["麻醉时长 (anesthesia_duration_min)", "麻醉总时长（分钟）", "麻醉记录原始值", "临床基础特征（与 Table 3 对齐）"],
            ["log(CRP) (log_crp)", "C 反应蛋白的自然对数", "CRP 取 ln；Table 3 中效应最强的炎症指标", "炎症代表特征（精简后仅保留此项）"],
            ["Shannon指数 (shannon)", "α 多样性（属水平）", "基于真实 DADA2 属水平相对丰度计算", "菌群代表特征（精简后仅保留此项）"],
        ],
    )
    _note(
        doc,
        "方案第3.2(B) 明确：临床=ASA+麻醉时长；炎症=log(CRP)；菌群=Shannon；合计 4 个特征，特征/样本比≈1:8。"
        "因此旧版炎症组的 WBC、log(PCT)、NLR 及旧版临床组的 age/sex/BMI/手术时长/阿片当量均不再进入模型。",
    )

    _h(doc, "0.2 模型定义（2×2 析因矩阵）", 2)
    _table(
        doc,
        ["", "不含菌群", "含菌群（+Shannon）"],
        [
            ["不含炎症", "Model A（临床）= ASA + 麻醉时长", "Model C（临床+菌群）= ASA + 麻醉时长 + Shannon"],
            ["含炎症（+logCRP）", "Model B（临床+炎症）= ASA + 麻醉时长 + log(CRP)", "Model E（全部）= ASA + 麻醉时长 + log(CRP) + Shannon"],
        ],
    )
    _note(
        doc,
        "析因分解：ΔAUC(B−A)=炎症单独贡献；ΔAUC(C−A)=菌群单独贡献；ΔAUC(E−B)=菌群在炎症基础上的增量（核心问题）；"
        "ΔAUC(E−C)=炎症在菌群基础上的增量；交互效应=(E−B)−(C−A)。",
    )

    _h(doc, "0.3 与旧版报告的差异（重要）", 2)
    _p(
        doc,
        "《优化分析方案.docx》第一部分与 Table 2/Table 4 模板中引用的 Model B AUC=0.757、B−A=+0.368「已确认显著」等，"
        "均来自旧版《AICU_output表格对照解析报告.docx》的富特征模型（临床 7 项 + 炎症 4 项 = 11 特征）。"
        "本版严格执行方案第3.2(B) 的特征精简后，同名模型的特征集与数值均已改变，二者不可直接混用。",
    )
    _table(
        doc,
        ["模型", "旧版特征集（0.757 报告）", "旧版 AUC", "本版特征集（精简）", "本版 AUC", "差异原因"],
        [
            [
                "Model A（临床）",
                "age, sex, BMI, ASA, 手术时长, 麻醉时长, 阿片当量（7）",
                "0.389",
                "ASA, 麻醉时长（2）",
                "0.472",
                "临床特征 7→2",
            ],
            [
                "Model B（临床+炎症）",
                "临床7 + WBC, log(CRP), log(PCT), NLR（11）",
                "0.757",
                "ASA, 麻醉时长, log(CRP)（3）",
                "0.576",
                "炎症仅留 log(CRP)，删去 WBC 等；WBC 为唯一组间显著项(P=0.004)",
            ],
            [
                "Model C（+菌群）",
                "临床7 + 炎症4 + 菌群多项（含炎症）",
                "0.677",
                "ASA, 麻醉时长, Shannon（不含炎症）",
                "0.510",
                "定义变更：2×2 析因下 C 不再含炎症",
            ],
            [
                "Model E（全部）",
                "（旧版无此模型；最接近旧 Model C）",
                "≈0.677",
                "ASA, 麻醉时长, log(CRP), Shannon（4）",
                "0.580",
                "新增；旧 Model C 的可比对象",
            ],
        ],
    )
    _note(
        doc,
        "关键提示：① 旧版「Model C=0.677」对应本版「Model E=0.580」，而本版「Model C=0.510」是全新定义（临床+菌群、不含炎症），"
        "跨文档比较时切勿按同名对应。② 本版 Model B AUC 下降主要因删去 WBC——它是 Table 1 中唯一组间显著的炎症指标(P=0.004)，"
        "而保留的 log(CRP) 组间不显著(P=0.143)；故「炎症预测价值有限」这一结论应限定为「仅以 log(CRP) 代表炎症」的前提下成立。",
    )

    doc.add_page_break()

    # 一、总体判断（基于本版精简模型结果，而非旧版 AUC=0.757）
    _h(doc, "一、基于本版结果的总体判断", 1)
    _p(
        doc,
        "研究定位（优化方案调整后）：入室炎症状态（尤其 CRP）可能与拔管延迟相关；"
        "本队列系统性检验了呼吸道菌群在组间差异、相关性、预测增益与交互效应等层面与炎症及预后的关系。"
        "在 n=34 精简特征模型下，菌群未展现独立于炎症的稳健增量增益；结果定位为假说生成型（hypothesis-generating）。",
    )
    a_auc = b_auc = e_auc = "—"
    if t2 is not None and not t2.empty:
        col_m = "模型" if "模型" in t2.columns else "Model"
        col_a = "AUC"
        for _, row in t2.iterrows():
            m = str(row[col_m])
            if "Model A" in m:
                a_auc = _fmt(row[col_a])
            elif "Model B" in m:
                b_auc = _fmt(row[col_a])
            elif "Model E" in m:
                e_auc = _fmt(row[col_a])
    eb = "—"
    if t2_fac is not None and not t2_fac.empty:
        hit = t2_fac[t2_fac["比较"].astype(str).str.contains("E - B", na=False)]
        if not hit.empty:
            eb = str(hit.iloc[0]["ΔAUC"])

    _table(
        doc,
        ["类别", "本版结论"],
        [
            [
                "核心发现（调整后）",
                f"精简 Model B（ASA+麻醉+logCRP）拔管延迟 AUC≈{b_auc}，置换不显著；"
                f"Model E≈{e_auc}；E−B 增量≈{eb}（近乎为 0）。"
                "不可再沿用旧版「Model B AUC=0.757」表述。",
            ],
            [
                "阴性/反常",
                "α/β/属水平菌群组间差异未达显著；Shannon–CRP 相关弱；中介不显著；"
                "WBC Early>Delayed（P=0.004）；ICU 滞留无组间差异；不良反应仅 5 例。",
            ],
            [
                "数据核查",
                f"Model A AUC≈{a_auc}（接近机会水平）；NLR 为 Excel 比值列直接解析，无法绝对值重算；"
                "COPD 零方差已自回归剔除。",
            ],
            [
                "Model C′",
                (
                    str(cprime.iloc[0]["reason"])
                    if cprime is not None and not cprime.empty
                    else "未评估"
                ),
            ],
        ],
    )

    # 二、Table 1
    doc.add_page_break()
    _h(doc, "二、Table 1  基线特征", 1)
    _p(doc, "对应优化方案：保持现有结构，补充脚注。")
    if t1 is not None:
        _table(doc, list(t1.columns), _df_rows(t1))
    if t1_notes is not None and not t1_notes.empty:
        _h(doc, "脚注", 2)
        for _, r in t1_notes.iterrows():
            _p(doc, f"{r['编号']} {r['内容']}")

    # Table 1-补充
    _h(doc, "三、Table 1-补充  两组菌群多样性比较", 1)
    _p(doc, "Early vs Delayed：Shannon / Chao1 / Pielou's J / Simpson；统一采用 Mann-Whitney U + Cliff's δ（优化方案第3.2(A)）。")
    if t1_sup is not None:
        _table(doc, list(t1_sup.columns), _df_rows(t1_sup))
    _note(doc, "解读：各组间 P 均 >0.05，效应量接近 0，与 Table 3 中 Shannon 回归不显著一致。")

    # β 多样性
    _h(doc, "四、β 多样性（PERMANOVA + betadisper）", 1)
    _p(doc, "Bray-Curtis；999 次置换；协变量调整：ASA + 麻醉时长。对应 Figure 3 / 优化方案 Figure X。")
    if perm is not None:
        show = perm.copy()
        for c in ("F", "R2", "p"):
            if c in show.columns:
                show[c] = show[c].map(lambda x: _fmt(x, 3 if c != "p" else 3))
        _table(doc, list(show.columns), _df_rows(show))
    if beta is not None and not beta.empty:
        _p(doc, f"betadisper：P={_fmt(beta.iloc[0]['p'])}（组内离散度检验；显著时需谨慎解读 PERMANOVA）。")

    # 属水平
    _h(doc, "五、Table X  属水平差异丰度（Wilcoxon + FDR）", 1)
    _p(doc, "首选方法；按 P 值排序。优化方案建议呈现 FDR<0.1 或提示性发现；LEfSe 作交叉验证。")
    if genus is not None and not genus.empty:
        g = genus.copy()
        # 优先展示 FDR<0.1；若无则展示 Top10
        if "q" in g.columns and (g["q"] < 0.1).any():
            g = g[g["q"] < 0.1].copy()
            _note(doc, "下列为 FDR<0.1 的属。")
        else:
            g = g.head(10).copy()
            _note(doc, "本版无 FDR<0.1 的具名属；下表为 Wilcoxon P 值最小的 Top 10（提示性，不作为阳性发现）。")
        cols = [c for c in ["genus", "early_mean_pct", "delayed_mean_pct", "p", "q", "enriched_group"] if c in g.columns]
        g2 = g[cols].rename(
            columns={
                "genus": "菌属",
                "early_mean_pct": "早期组丰度(%)",
                "delayed_mean_pct": "延迟组丰度(%)",
                "p": "Wilcoxon P",
                "q": "FDR-P",
                "enriched_group": "富集方向",
            }
        )
        for c in ("早期组丰度(%)", "延迟组丰度(%)"):
            if c in g2.columns:
                g2[c] = g2[c].map(lambda x: _fmt(x, 3))
        for c in ("Wilcoxon P", "FDR-P"):
            if c in g2.columns:
                g2[c] = g2[c].map(lambda x: _fmt(x, 3))
        _table(doc, list(g2.columns), _df_rows(g2))
    if lefse is not None and not lefse.empty:
        _p(doc, "LEfSe 交叉验证（摘要）：")
        cols = [c for c in ["genus", "p", "q", "lda", "enriched_group"] if c in lefse.columns]
        L = lefse[cols].head(5).copy()
        _table(doc, list(L.columns), _df_rows(L))
        _note(doc, "真实数据中优势信号为 Unknown；Model C′ 未触发（见文末核查表）。")

    # Table 2 主表
    doc.add_page_break()
    _h(doc, "六、Table 2（更新版）  LOO-CV 模型性能（A/B/C/E）", 1)
    _p(doc, "主分析结局：拔管延迟（二分类）。特征精简：临床=ASA+麻醉时长；炎症=log(CRP)；菌群=Shannon。不良反应/MLP 见补充材料。")
    if t2 is not None:
        show = t2.copy()
        # 统一列名
        rename = {
            "Target": "预测结局",
            "Model": "模型",
            "Accuracy": "准确率",
            "Sensitivity": "灵敏度",
            "Specificity": "特异度",
            "AUC_CI_low": "AUC CI下限",
            "AUC_CI_high": "AUC CI上限",
            "Permutation_P": "置换检验P",
        }
        show = show.rename(columns={k: v for k, v in rename.items() if k in show.columns})
        prefer = ["模型", "AUC", "AUC CI下限", "AUC CI上限", "灵敏度", "特异度", "F1", "准确率", "置换检验P"]
        cols = [c for c in prefer if c in show.columns]
        _table(doc, cols, _df_rows(show[cols]))
    _note(
        doc,
        "说明：本版 Model B/E 置换 P>0.05，AUC≈0.58；勿与优化方案起草时引用的旧版 Model B（AUC=0.757）混用。",
    )

    # 连续变量敏感性
    _h(doc, "七、Table 2 敏感性  拔管时间连续变量（log）", 1)
    _p(doc, "模型：log(拔管时间) ~ Shannon + log(CRP) + ASA + 麻醉时长；系数为标准化回归系数(β)，可横向比较效应大小（优化方案第3.2(C)）。")
    if t2_cont is not None:
        _table(doc, list(t2_cont.columns), _df_rows(t2_cont))
    _note(doc, "全部协变量 P>0.05，R²≈0.13；标准化后麻醉时长 |β| 最大、log(CRP) 与 Shannon 较小，与二分类模型弱信号一致。")

    # 析因
    _h(doc, "八、Table 2-附  析因分解（ΔAUC）", 1)
    _p(doc, "核心问题：E−B = 菌群在炎症基础上的增量。")
    if t2_fac is not None:
        rows = []
        for _, r in t2_fac.iterrows():
            label = str(r["比较"])
            if "E - B" in label:
                interpret = "核心：菌群在炎症上的增量≈0，无协同增益"
            elif "B - A" in label:
                interpret = "炎症单独贡献有限（本版精简模型下未达显著）"
            elif "C - A" in label:
                interpret = "菌群单独贡献弱"
            elif "E - C" in label:
                interpret = "炎症在菌群基础上的增量有限"
            elif "交互" in label:
                interpret = "负值提示无协同、接近独立/拮抗，样本量下不稳定"
            else:
                interpret = "—"
            rows.append([label, str(r["ΔAUC"]), str(r.get("95%CI", "—")), str(r.get("Bootstrap_P", "—")), interpret])
        _table(doc, ["比较", "ΔAUC", "95%CI（Bootstrap）", "P", "解读"], rows)

    # Table 3
    doc.add_page_break()
    _h(doc, "九、Table 3  多因素 Logistic 回归", 1)
    _p(doc, "控制 ASA、麻醉时长；评估 Shannon 与 log(CRP)。优化方案强调 log(CRP) 的提示性关联（suggestive association）。")
    if t3 is not None:
        show = t3.copy()
        show["变量"] = show["variable"].map(lambda v: VAR_CN.get(v, v))
        show["OR(95%CI)"] = show.apply(
            lambda r: f"{r['OR']:.2f}（{r['CI_low']:.2f}–{r['CI_high']:.2f}）", axis=1
        )
        cols_out = ["outcome_label", "model", "变量", "OR(95%CI)", "P_fmt", "N"]
        rename = {"outcome_label": "结局", "model": "模型", "P_fmt": "P值"}
        show = show[cols_out].rename(columns=rename)
        _table(doc, list(show.columns), _df_rows(show))
    _note(doc, "拔管延迟结局中 log(CRP) OR≈1.62–1.64，P≈0.09–0.11，为边缘趋势；Shannon 不显著。")

    # 相关矩阵
    _h(doc, "十、菌群多样性 × 炎症指标相关（Spearman + FDR）", 1)
    if corr is not None:
        show = corr.copy()
        if "row" in show.columns:
            show = show.rename(columns={"row": "多样性", "col": "炎症指标", "rho": "ρ", "p": "P", "q": "FDR-q"})
        for c in ("ρ", "P", "FDR-q"):
            if c in show.columns:
                show[c] = show[c].map(lambda x: _fmt(x, 3))
        _table(doc, list(show.columns), _df_rows(show))
    _note(doc, "Shannon–CRP ρ≈0.08，|ρ|<0.3 → 中介分析按优化方案作简短阴性报告。")

    # 中介
    _h(doc, "十一、中介分析（Shannon → log(CRP) → 拔管时间）", 1)
    if med is not None:
        if med.shape[1] == 2 and med.columns[0] != "path":
            d = dict(zip(med.iloc[:, 0].astype(str), med.iloc[:, 1]))
        else:
            d = med.set_index(med.columns[0]).iloc[:, 0].to_dict() if med.shape[1] >= 2 else {}
        rows = [
            ["总效应 c", _fmt(d.get("c_total"))],
            ["直接效应 c'", _fmt(d.get("c_direct"))],
            ["间接效应 ab", _fmt(d.get("ab_indirect"))],
            ["ab 95%CI", f"{_fmt(d.get('ab_ci_low'))} – {_fmt(d.get('ab_ci_high'))}"],
            ["中介比例", _fmt(float(d.get("mediation_proportion", float("nan"))) * 100, 1) + "%" if d.get("mediation_proportion") is not None else "—"],
            ["间接效应显著", str(d.get("indirect_significant", "—"))],
        ]
        _table(doc, ["路径", "估计值"], rows)
    _p(doc, "结论：未观察到显著中介效应（间接效应 CI 含 0）；不作展开机制解读。")

    # 补充材料
    doc.add_page_break()
    _h(doc, "十二、补充材料", 1)
    _h(doc, "12.1 不良反应描述性（不做假设检验）", 2)
    if ae is not None:
        _table(doc, list(ae.columns), _df_rows(ae))
    _note(doc, "因事件数过少(n=5)，不具备可靠统计学意义，仅供参考。")

    _h(doc, "12.2 Model D（MLP）及不良反应模型（补充）", 2)
    if t2_supp is not None and not t2_supp.empty:
        show = t2_supp.copy()
        rename = {
            "Target": "预测结局",
            "Model": "模型",
            "AUC_CI_low": "AUC CI下限",
            "AUC_CI_high": "AUC CI上限",
            "Permutation_P": "置换检验P",
        }
        show = show.rename(columns={k: v for k, v in rename.items() if k in show.columns})
        # 只保留关键列
        cols = [c for c in ["预测结局", "模型", "AUC", "AUC CI下限", "AUC CI上限", "置换检验P", "Sensitivity", "灵敏度"] if c in show.columns]
        for c in ("AUC", "AUC CI下限", "AUC CI上限"):
            if c in show.columns:
                show[c] = show[c].map(lambda x: _fmt(x, 3))
        _table(doc, cols, _df_rows(show[cols]))
    _note(doc, "探索性/概念验证分析，样本量限制下结果解释需谨慎。")

    _h(doc, "12.3 NLR 口径核查", 2)
    if nlr is not None:
        _table(doc, list(nlr.columns), _df_rows(nlr))

    _h(doc, "12.4 Model C′ 触发评估", 2)
    if cprime is not None:
        _table(doc, list(cprime.columns), _df_rows(cprime))

    _h(doc, "12.5 敏感性分析（QC 队列）", 2)
    if sens is not None:
        show = sens.copy()
        labels = {"strict": "严格QC", "main": "主分析", "relaxed": "放宽QC"}
        if "cohort" in show.columns:
            show["cohort"] = show["cohort"].map(lambda c: labels.get(c, c))
        _table(doc, list(show.columns), _df_rows(show))

    # 叙事建议
    _h(doc, "十三、论文叙事建议（对应优化方案情景一）", 1)
    _p(
        doc,
        "本版结果最接近优化方案「情景一」：Model E 与 Model B 接近，菌群组间差异有限。"
        "建议表述：炎症指标（logCRP）呈现边缘提示性关联；菌群在多样性、组成与预测层面均未提供稳健独立或增量信息；"
        "定位为小样本假说生成型研究，讨论测序深度、注释 Incomplete（Unknown/Unassigned）与样本量限制。",
    )
    _p(doc, "Results 建议顺序：Table1 → 菌群α/β/属 → 相关 → Table2+析因 → Table3 → 中介阴性简报 → 补充材料。")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("— 文档结束 —")
    _font(fr, size=Pt(9))
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="按优化分析方案整合 output 表格为 Word")
    ap.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    ap.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    args = ap.parse_args()
    path = build(args.output_dir, args.docx)
    print(f"已生成: {path}")
