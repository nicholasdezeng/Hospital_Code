#!/usr/bin/env python3
"""生成 Table 1–3 具体解析 Word 文档（仅表格内容）。"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = Path("/Users/zengyinghao/Downloads/output")
DEFAULT_DOCX = ROOT.parent / "ALL" / "AICU_output表格对照解析报告.docx"

VAR_CN = {
    "shannon": "Shannon指数",
    "log_crp": "log(CRP)",
    "asa": "ASA分级",
    "anesthesia_duration_min": "麻醉时长(min)",
}


def _set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    r = style._element.get_or_add_rPr()
    r.rFonts.set(qn("w:eastAsia"), "宋体")


def _set_run_font(run, *, east_asia: str = "宋体", size: Pt | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    r = run._element.get_or_add_rPr()
    r.rFonts.set(qn("w:eastAsia"), east_asia)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run)


def _para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=Pt(11))


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def _read_csv(path: Path) -> pd.DataFrame | None:
    if not path.exists():
        return None
    return pd.read_csv(path, encoding="utf-8-sig")


def _first_csv(*paths: Path) -> pd.DataFrame | None:
    for path in paths:
        df = _read_csv(path)
        if df is not None:
            return df
    return None


def _parse_table2(table2: pd.DataFrame) -> list[str]:
    target_col = "Target" if "Target" in table2.columns else "预测结局"
    model_col = "Model" if "Model" in table2.columns else "模型"
    notes: list[str] = []
    for target_key, target_label in [("Extubation|delay", "拔管延迟"), ("Adverse|不良", "不良反应")]:
        sub = table2[table2[target_col].astype(str).str.contains(target_key, case=False, na=False)]
        if sub.empty:
            continue
        notes.append(f"【{target_label}】")
        best_idx = sub["AUC"].astype(float).idxmax()
        for _, r in sub.iterrows():
            model = r[model_col]
            auc = float(r["AUC"])
            perm = r.get("Permutation_P", r.get("置换检验P", ""))
            ci_l = r.get("AUC_CI_low", r.get("AUC CI下限", ""))
            ci_h = r.get("AUC_CI_high", r.get("AUC CI上限", ""))
            acc = r.get("Accuracy", r.get("准确率", ""))
            sens = r.get("Sensitivity", r.get("灵敏度", ""))
            spec = r.get("Specificity", r.get("特异度", ""))
            f1 = r.get("F1", "")
            tag = "（本结局 AUC 最高）" if r.name == best_idx else ""
            notes.append(
                f"  {model}：AUC={auc:.3f}（95%CI {ci_l}–{ci_h}），"
                f"准确率={acc}，灵敏度={sens}，特异度={spec}，F1={f1}，置换检验 P={perm}{tag}。"
            )
        if target_label == "拔管延迟":
            b = sub[sub[model_col].astype(str).str.contains("Model B", na=False)]
            c = sub[sub[model_col].astype(str).str.contains("Model C", na=False)]
            if not b.empty and not c.empty:
                b_auc, c_auc = float(b.iloc[0]["AUC"]), float(c.iloc[0]["AUC"])
                if b_auc > c_auc:
                    notes.append(
                        f"  解析：Model B（+炎症）AUC={b_auc:.3f} 高于 Model C（+菌群）AUC={c_auc:.3f}，"
                        "加入菌群特征未提升判别能力。"
                    )
                else:
                    notes.append(
                        f"  解析：Model C AUC={c_auc:.3f} 高于 Model B AUC={b_auc:.3f}，"
                        "菌群特征有增量价值。"
                    )
        if target_label == "不良反应":
            notes.append("  解析：不良反应事件数少（Table 1 共 5 例），各模型 AUC 多接近 0.5，结果需谨慎解读。")
    return notes


def build_document(output_dir: Path, docx_path: Path) -> Path:
    output_dir = output_dir.resolve()
    tab_dir = output_dir / "tables"
    fig_dir = output_dir / "figures"
    summary = json.loads((output_dir / "run_summary.json").read_text(encoding="utf-8"))

    table1 = _read_csv(tab_dir / "table1_baseline.csv")
    table2 = _first_csv(tab_dir / "table2_model_performance.csv", fig_dir / "table2_model_performance.csv")
    table3 = _read_csv(tab_dir / "table3_multivariable_or.csv")

    doc = Document()
    _set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AICU output 表格具体解析\n（Table 1–3）")
    _set_run_font(tr, east_asia="黑体", size=Pt(16), bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"n={summary.get('n_samples', '—')}（Early {summary.get('n_early', '—')} / "
        f"Delayed {summary.get('n_delayed', '—')}）｜"
        f"Excel：04/05/06 工作表｜{datetime.now().strftime('%Y-%m-%d')}"
    )
    _set_run_font(sr, size=Pt(10))

    # ── Table 1 ──
    doc.add_page_break()
    _heading(doc, "Table 1  基线特征（Excel：04_基线特征）", 1)
    _para(doc, "方案要求：按 Early / Delayed 拔管组比较基线变量，报告 P 值。")

    if table1 is not None:
        cols = list(table1.columns)
        rows = [[str(r[c]) for c in cols] for _, r in table1.iterrows()]
        _add_table(doc, cols, rows)

    # ── Table 2 ──
    doc.add_page_break()
    _heading(doc, "Table 2  模型性能（Excel：05_模型性能）", 1)
    _para(doc, "方案要求：LOO-CV 下 Model A/B/C（及 MLP）对拔管延迟、不良反应的 AUC 及性能指标。")

    if table2 is not None:
        # 统一列名展示
        show = table2.copy()
        rename = {
            "Target": "预测结局", "Model": "模型", "Accuracy": "准确率",
            "Sensitivity": "灵敏度", "Specificity": "特异度",
            "AUC_CI_low": "AUC CI下限", "AUC_CI_high": "AUC CI上限",
            "Permutation_P": "置换检验P",
        }
        show = show.rename(columns={k: v for k, v in rename.items() if k in show.columns})
        for col in ("AUC", "AUC CI下限", "AUC CI上限", "准确率", "灵敏度", "特异度", "F1"):
            if col in show.columns:
                show[col] = show[col].map(lambda x: f"{float(x):.3f}" if pd.notna(x) else "")
        if "置换检验P" in show.columns:
            show["置换检验P"] = show["置换检验P"].map(
                lambda x: f"{float(x):.3f}" if pd.notna(x) else ""
            )
        cols = list(show.columns)
        rows = [[str(r[c]) for c in cols] for _, r in show.iterrows()]
        _add_table(doc, cols, rows)
        _heading(doc, "逐模型解析", 2)
        for line in _parse_table2(table2):
            _para(doc, line)

    # ── Table 3 ──
    doc.add_page_break()
    _heading(doc, "Table 3  多因素 Logistic 回归（Excel：06_多因素回归）", 1)
    _para(doc, "方案要求：控制 ASA、麻醉时长，评估 Shannon 与 log(CRP) 对结局的独立效应。")

    if table3 is not None:
        show = table3.copy()
        show["变量"] = show["variable"].map(lambda v: VAR_CN.get(v, v))
        show["OR(95%CI)"] = show.apply(
            lambda r: f"{r['OR']:.2f}（{r['CI_low']:.2f}–{r['CI_high']:.2f}）", axis=1
        )
        cols = ["outcome_label", "model", "变量", "OR(95%CI)", "P_fmt", "N"]
        show_cols = ["结局", "模型", "变量", "OR(95%CI)", "P值", "N"]
        rows = []
        for _, r in show.iterrows():
            rows.append([str(r["outcome_label"]), str(r["model"]), str(r["变量"]),
                         str(r["OR(95%CI)"]), str(r["P_fmt"]), str(int(r["N"]))])
        _add_table(doc, show_cols, rows)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="生成 Table 1–3 具体解析 Word")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    args = parser.parse_args()
    print(f"已生成: {build_document(args.output_dir, args.docx)}")
